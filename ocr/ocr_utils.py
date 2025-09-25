import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import re

import pdfplumber
import PyPDF2
from PIL import Image
import pytesseract
from docx import Document

from config import OCR_DPI, OCR_LANGUAGES, TESSERACT_CONFIG

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class OCRProcessor:
    """OCR processor for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a file based on its format
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: Contains 'text', 'pages', 'metadata' keys
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_ext == '.txt':
                return self._extract_from_txt(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg']:
                return self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'text': '',
                'pages': 0,
                'metadata': {'error': str(e)},
                'success': False
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF using pdfplumber with fallback to PyPDF2"""
        text_content = []
        pages = 0
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                pages = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        
            if not text_content:
                # Fallback to PyPDF2
                logger.info(f"pdfplumber failed, trying PyPDF2 for {file_path}")
                text_content = self._extract_with_pypdf2(file_path)
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {str(e)}")
            text_content = self._extract_with_pypdf2(file_path)
        
        full_text = '\n\n'.join(text_content) if text_content else ''
        
        return {
            'text': full_text,
            'pages': pages,
            'metadata': {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'extraction_method': 'pdfplumber'
            },
            'success': len(full_text.strip()) > 0
        }
    
    def _extract_with_pypdf2(self, file_path: Path) -> list:
        """Fallback PDF extraction with PyPDF2"""
        text_content = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {str(e)}")
        
        return text_content
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n\n'.join(text_content)
            
            return {
                'text': full_text,
                'pages': 1,  # DOCX doesn't have pages concept
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'extraction_method': 'python-docx'
                },
                'success': len(full_text.strip()) > 0
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return {
                'text': '',
                'pages': 0,
                'metadata': {'error': str(e)},
                'success': False
            }
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                
            return {
                'text': text,
                'pages': 1,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'extraction_method': 'direct_read'
                },
                'success': len(text.strip()) > 0
            }
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            return {
                'text': '',
                'pages': 0,
                'metadata': {'error': str(e)},
                'success': False
            }
    
    def _extract_from_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from images using Tesseract OCR with Malayalam support"""
        try:
            # Open and process image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Detect script type for better OCR configuration
            script_type = self._detect_image_script_type(image)
            
            # Choose appropriate Tesseract configuration
            if script_type == 'malayalam':
                config = TESSERACT_CONFIG.get('malayalam', '--psm 6')
                languages = 'mal+eng'  # Malayalam + English
            elif script_type == 'mixed':
                config = TESSERACT_CONFIG.get('mixed', '--psm 6')
                languages = 'mal+eng'  # Both languages
            else:
                config = TESSERACT_CONFIG.get('english', '--psm 6') 
                languages = 'eng'  # English only
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image, config=config, lang=languages)
            
            # Post-process extracted text
            text = self._post_process_ocr_text(text, script_type)
            
            return {
                'text': text,
                'pages': 1,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'extraction_method': 'tesseract_ocr_multilingual',
                    'image_size': image.size,
                    'detected_script': script_type,
                    'ocr_languages': languages,
                    'ocr_config': config
                },
                'success': len(text.strip()) > 0
            }
            
        except Exception as e:
            logger.error(f"Image OCR extraction failed: {str(e)}")
            return {
                'text': '',
                'pages': 0,
                'metadata': {'error': str(e)},
                'success': False
            }
    
    def _detect_image_script_type(self, image: Image.Image) -> str:
        """
        Detect the primary script type in an image for better OCR
        This is a simplified approach - could be enhanced with ML models
        """
        try:
            # Quick OCR with English only to check for Malayalam presence
            sample_text = pytesseract.image_to_string(image, config='--psm 6', lang='eng')
            
            # Check for Malayalam Unicode characters that might have been detected
            malayalam_pattern = r'[\u0D00-\u0D7F]'
            
            # Also try a small Malayalam OCR sample if Tesseract supports it
            try:
                malayalam_sample = pytesseract.image_to_string(image, config='--psm 6', lang='mal')
                malayalam_chars = len(re.findall(malayalam_pattern, malayalam_sample))
                total_chars = len(re.sub(r'\s+', '', malayalam_sample))
                
                if total_chars > 0 and malayalam_chars / total_chars > 0.1:
                    return 'malayalam' if malayalam_chars / total_chars > 0.7 else 'mixed'
                    
            except Exception:
                # Malayalam language pack not available, use heuristics
                pass
            
            # Fallback: check if the English OCR result suggests non-Latin content
            if len(sample_text.strip()) < 10 or len([c for c in sample_text if c.isalnum()]) < len(sample_text) * 0.3:
                return 'mixed'  # Likely contains non-English script
            
            return 'english'
            
        except Exception as e:
            logger.warning(f"Script detection failed: {str(e)}")
            return 'mixed'  # Default to mixed for safety
    
    def _post_process_ocr_text(self, text: str, script_type: str) -> str:
        """Post-process OCR text based on detected script type"""
        if not text:
            return text
        
        # Common OCR error corrections
        text = re.sub(r'\n{3,}', '\n\n', text)  # Reduce multiple newlines
        text = re.sub(r'[ \t]{2,}', ' ', text)  # Reduce multiple spaces
        
        if script_type == 'malayalam' or script_type == 'mixed':
            # Malayalam-specific corrections
            # Fix common OCR errors for Malayalam
            malayalam_corrections = {
                'ന്‍': 'ന്',  # Combining marks
                'ര്‍': 'ര്',
                'ല്‍': 'ल्',
                '്‌': '്',  # Virama corrections
            }
            
            for wrong, correct in malayalam_corrections.items():
                text = text.replace(wrong, correct)
        
        # Clean up extra whitespace
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        
        return '\n'.join(cleaned_lines)

def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Convenience function to extract text from any supported file
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        dict: Extraction result with text, metadata, and success status
    """
    processor = OCRProcessor()
    return processor.extract_text(file_path)